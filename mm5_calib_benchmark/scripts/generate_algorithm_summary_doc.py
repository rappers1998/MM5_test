from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_ROOT = PROJECT_ROOT / "mm5_calib_benchmark" / "outputs" / "mm5_benchmark"
RUNS_ROOT = PROJECT_ROOT / "runs"
TODAY = datetime.now().strftime("%Y%m%d")
OUTPUT_DOC = RUNS_ROOT / f"MM5_校准算法总结_中文版_{TODAY}.docx"


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _safe_rel(path: Path) -> str:
    try:
        return str(path.relative_to(PROJECT_ROOT))
    except ValueError:
        return str(path)


def _fmt_float(value: Any, digits: int = 4) -> str:
    if value is None:
        return "N/A"
    try:
        value_f = float(value)
    except (TypeError, ValueError):
        return str(value)
    if value_f != value_f:
        return "N/A"
    return f"{value_f:.{digits}f}"


def _summary_row(output_name: str) -> dict[str, Any]:
    return _read_json(OUTPUT_ROOT / output_name / "metrics" / "summary.json")


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    run.append(run_props)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_code_block(document: Document, text: str) -> None:
    for line in text.rstrip().splitlines():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        paragraph.paragraph_format.space_after = Pt(0)


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _add_image(document: Document, path: Path, caption: str, width_in: float = 6.4) -> None:
    if not path.exists():
        document.add_paragraph(f"[Image missing] {_safe_rel(path)}")
        return
    picture_par = document.add_paragraph()
    picture_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_par.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = document.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


METHODS: list[dict[str, Any]] = [
    {
        "key": "m0",
        "display": "M0_MM5_OFFICIAL",
        "track": "thermal",
        "available_tracks": ["thermal", "uv"],
        "output_name": "method_M0_mm5_official_thermal",
        "image_path": OUTPUT_ROOT / "method_M0_mm5_official_thermal" / "viz" / "qualitative_grid_best_mid_worst.png",
        "code_paths": [
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m0_mm5_official" / "run.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "alignment.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "board.py",
        ],
        "family": "官方 stereo + 平面单应 + scene-level affine tune",
        "objective": "作为最轻量的基线，验证 MM5 官方标定在当前 raw RGB→thermal 任务上的可迁移性。",
        "implementation_note": "当前工作区版本不是“纯官方输出”，而是在官方 stereo 之后又增加了固定平面深度单应变换和 scene-level tune。",
        "workflow": [
            "直接读取 MM5 官方 stereo 标定。",
            "在固定平面深度 700 mm 下计算 RGB 到目标模态的 homography。",
            "把 RGB 图像与 RGB 语义 mask 一起投影到目标模态。",
            "用目标边缘距离、互信息和 coverage 稳定性做 coarse-to-fine scene tuning。",
        ],
        "strengths": [
            "实现最简单，适合做 sanity check 和回归基线。",
            "可以快速暴露官方标定与当前数据分布之间的偏差。",
        ],
        "limitations": [
            "强依赖平面假设，遇到明显景深或物体起伏时会退化。",
            "不利用深度，因此边界和遮挡处理能力最弱。",
        ],
        "references": [
            ("MM5 数据集与官方 calibration 来源", "https://doi.org/10.1016/j.inffus.2025.103516"),
        ],
    },
    {
        "key": "m1",
        "display": "M1_ZHANG_OPENCV",
        "track": "thermal",
        "available_tracks": ["thermal", "uv"],
        "output_name": "method_M1_zhang_opencv_thermal",
        "image_path": OUTPUT_ROOT / "method_M1_zhang_opencv_thermal" / "viz" / "qualitative_grid_best_mid_worst.png",
        "code_paths": [
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m1_zhang" / "run.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "board.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "alignment.py",
        ],
        "family": "棋盘格 stereo 标定 + 平面单应 + scene tune",
        "objective": "在同一 benchmark 下测试传统棋盘格 stereo 标定是否优于直接使用官方外参。",
        "implementation_note": "代码先用 calibration 目录下的 board observation 做 stereo 标定；若失败，则回退到官方 stereo。",
        "workflow": [
            "读取 calibration 中的多帧棋盘格观测。",
            "依据 Zhang 标定思路估计两相机 stereo 外参。",
            "用固定平面深度执行 homography warp。",
            "再进行较小搜索半径的 scene-level affine tune。",
        ],
        "strengths": [
            "比 M0 更贴近当前装置的实际 board 观测。",
            "实现稳定，可同时服务 thermal 和 UV 方向。",
        ],
        "limitations": [
            "推理阶段仍然是平面 homography，因此并不能处理真实 3D 视差。",
            "对跨模态局部错位的修正能力有限。",
        ],
        "references": [
            ("Zhang 2000: A flexible new technique for camera calibration", "https://doi.org/10.1109/34.888718"),
            ("MM5 数据集论文", "https://doi.org/10.1016/j.inffus.2025.103516"),
        ],
    },
    {
        "key": "m2",
        "display": "M2_SU2025_XOFTR_FALLBACK",
        "track": "thermal",
        "available_tracks": ["thermal"],
        "output_name": "method_M2_su2025_xoftr_thermal",
        "image_path": OUTPUT_ROOT / "method_M2_su2025_xoftr_thermal" / "viz" / "qualitative_grid_best_mid_worst.png",
        "code_paths": [
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m2_su2025_xoftr" / "run.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "alignment.py",
        ],
        "family": "M1 stereo + 全局 homography + 跨模态局部特征仿射 refine + scene tune",
        "objective": "模拟跨模态特征匹配标定思路，测试在不依赖深度的情况下，局部特征是否能改善热像对齐。",
        "implementation_note": "方法名里有 XoFTR，但当前工作区实现并没有加载 transformer 权重，而是采用 SIFT/ORB 的 fallback 特征匹配，再做 partial affine refine。",
        "workflow": [
            "优先复用 M1 已保存 stereo，否则回退到官方 stereo。",
            "先做平面 homography 对齐。",
            "对 warped RGB 和 thermal 做 SIFT/ORB 特征匹配，并估计 partial affine。",
            "在 affine refine 后再做一次较窄搜索的 scene tuning。",
        ],
        "strengths": [
            "比纯 homography 多了一层局部几何修正。",
            "不需要深度图，适合做无深度条件下的中间强度基线。",
        ],
        "limitations": [
            "当前实现不是原论文的完整 cross-modal transformer 方案。",
            "跨模态特征不足时会自动放弃 feature refine。",
        ],
        "references": [
            ("Su et al. 2025: Joint Calibration Method of Thermal Infrared-Visible Based on Cross Modal Feature Matching", "https://doi.org/10.5194/isprs-annals-X-1-W2-2025-139-2025"),
            ("XoFTR: Cross-modal Feature Matching Transformer", "https://arxiv.org/abs/2404.09692"),
        ],
    },
    {
        "key": "m3",
        "display": "M3_JAY2025_REGISTRATION_FALLBACK",
        "track": "uv",
        "available_tracks": ["uv"],
        "output_name": "method_M3_jay2025_sgm_uv",
        "image_path": OUTPUT_ROOT / "method_M3_jay2025_sgm_uv" / "viz" / "qualitative_grid_best_mid_worst.png",
        "code_paths": [
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m3_jay2025_sgm" / "run.py",
        ],
        "family": "homography + dense optical flow fallback",
        "objective": "作为当前仓库里 UV 方向的 dense warp 对照组，近似验证“局部形变补偿”类方法是否能改善 registration。",
        "implementation_note": "当前实现并不是 Jay 2025 论文中的 SGM + disparity + pixel filling 流程，而是 homography 后接 Farneback dense optical flow 的 fallback 版本。",
        "workflow": [
            "读取 M1 stereo 或官方 stereo。",
            "先执行全局 homography warp。",
            "在目标模态上计算 Farneback dense optical flow。",
            "用 flow field 对 mask / source image / valid mask 做 dense remap。",
        ],
        "strengths": [
            "coverage 很高，说明 dense warp 容易把整张图映射满。",
            "适合作为“仅靠图像形变”的对照方法。",
        ],
        "limitations": [
            "当前仅有 UV benchmark 输出，没有 thermal 主图结果。",
            "不是论文原始 SGM 实现，因此只能视为 Jay 家族的 fallback baseline。",
        ],
        "references": [
            ("Jay et al. 2025: Registration of close-range, multi-lens multispectral imagery by retrieving the scene 3D structure", "https://doi.org/10.1016/j.isprsjprs.2025.06.001"),
            ("Farneback 2003: Two-Frame Motion Estimation Based on Polynomial Expansion", "https://doi.org/10.1007/3-540-45103-X_50"),
        ],
    },
    {
        "key": "m4",
        "display": "M4_DEPTHBRIDGE_ADAPTED",
        "track": "thermal",
        "available_tracks": ["thermal", "uv"],
        "output_name": "method_M4_muhovic_depthbridge_thermal",
        "image_path": OUTPUT_ROOT / "method_M4_muhovic_depthbridge_thermal" / "viz" / "qualitative_grid_best_mid_worst.png",
        "code_paths": [
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m4_muhovic_depthbridge" / "run.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "alignment.py",
        ],
        "family": "depth projection + z-buffer + support-region fill + scene tune",
        "objective": "把显式深度纳入投影，验证 depth bridge / parallax-aware 家族对热像标定的帮助。",
        "implementation_note": "当前实现参考了 parallax-aware annotation transfer 的思路，但使用的是 MM5 的 RGB-depth 数据，而不是论文中的 LiDAR 方案。",
        "workflow": [
            "优先复用 M5 或 M1 已保存 stereo。",
            "用 raw depth 把 RGB 像素反投影到 3D，再投影到目标模态。",
            "使用 z-buffer、support mask 和最近种子补洞。",
            "最后做一次 scene-level tune 修正整体位置和尺度。",
        ],
        "strengths": [
            "比纯 homography 更能处理视差和遮挡。",
            "是后续 M6 和 M7 的重要几何骨架。",
        ],
        "limitations": [
            "深度空洞和投影 support 稀疏时，补洞质量会成为瓶颈。",
            "当前不是原论文的一比一复现，而是针对 MM5 的 adapted 版本。",
        ],
        "references": [
            ("Muhovič & Perš 2023: Joint Calibration of a Multimodal Sensor System for Autonomous Vehicles", "https://doi.org/10.3390/s23125676"),
            ("作者公开实现", "https://github.com/JonNatanael/multimodal_calibration"),
        ],
    },
    {
        "key": "m5",
        "display": "M5_EPNP_BASELINE",
        "track": "thermal",
        "available_tracks": ["thermal", "uv"],
        "output_name": "method_M5_epnp_baseline_thermal",
        "image_path": OUTPUT_ROOT / "method_M5_epnp_baseline_thermal" / "viz" / "qualitative_grid_best_mid_worst.png",
        "code_paths": [
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m5_epnp" / "run.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "board.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "alignment.py",
        ],
        "family": "board stereo + EPnP relative pose + depth projection baseline",
        "objective": "给深度投影链路配一个更几何化的外参估计基线，作为当前非 MAR 系列的强基线之一。",
        "implementation_note": "calibrate() 会先从 board observation 求 stereo，再尝试用 EPnP 形式估计相对位姿；若失败则回退到基础 stereo。",
        "workflow": [
            "基于 board 数据求基础 stereo 标定。",
            "对外参进一步执行 EPnP 相对位姿估计。",
            "使用 raw depth 做 3D 投影与 tighter support fill。",
            "再做 scene tuning 获取最终结果。",
        ],
        "strengths": [
            "在 M7 出现前，它是当前 thermal 非 MAR 方法里表现最稳的一条线。",
            "比 M4 的投影更保守，补洞半径更小，边界通常更干净。",
        ],
        "limitations": [
            "仍然依赖单次 scene tuning 收敛到合适位置。",
            "没有边界级 refinement，复杂热边缘仍会偏软。",
        ],
        "references": [
            ("Lepetit et al. 2009: EPnP: An Accurate O(n) Solution to the PnP Problem", "https://doi.org/10.1007/s11263-008-0152-6"),
            ("EPFL / IRI EPnP 页面", "https://www.epfl.ch/labs/cvlab/software/multi-view-stereo/epnp/"),
        ],
    },
    {
        "key": "m6",
        "display": "M6_MAR_EDGE_REFINE",
        "track": "thermal",
        "available_tracks": ["thermal"],
        "output_name": "method_M6_mar_edge_refine_thermal",
        "image_path": OUTPUT_ROOT / "method_M6_mar_edge_refine_thermal" / "viz" / "qualitative_grid_best_mid_worst.png",
        "code_paths": [
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m6_mar_edge_refine" / "run.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "mar_edge_refine.py",
        ],
        "family": "M4 几何骨架 + scene tune + random-walker 边界 refine",
        "objective": "在不直接调用旧 MAR 历史结果的前提下，重建一条更接近 MAR 风格的边界优化链路。",
        "implementation_note": "它不是旧版完整 MAR full pipeline；它是 benchmark 内部的 MAR-style approximation，用于和历史 MAR 分开比较。",
        "workflow": [
            "先用 M4 风格的 depth bridge 结果作为 base。",
            "在 target 边缘图上执行 scene-level tune。",
            "使用 random walker 只在边界带内做重新标注。",
            "若面积比、binary IoU 或连通域数量超出 guard，则回退。",
        ],
        "strengths": [
            "比纯几何投影更关注边界质量。",
            "适合验证 MAR 类后处理对 contour 的增益。",
        ],
        "limitations": [
            "当前 mean_iou 并没有超过 M5。",
            "它不等于旧版完整 MAR 主流程，不能把它当成历史高分 MAR 的同义词。",
        ],
        "references": [
            ("MM5 论文中的 MAR 背景与数据来源", "https://doi.org/10.1016/j.inffus.2025.103516"),
            ("Grady 2006: Random Walks for Image Segmentation", "https://doi.org/10.1109/TPAMI.2006.233"),
        ],
    },
    {
        "key": "m7",
        "display": "M7_DEPTH_GUIDED_SELF_CAL_THERMAL",
        "track": "thermal",
        "available_tracks": ["thermal"],
        "output_name": "method_M7_depth_guided_selfcal_thermal",
        "image_path": OUTPUT_ROOT / "method_M7_depth_guided_selfcal_thermal" / "viz" / "qualitative_grid_best_mid_worst.png",
        "code_paths": [
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m7_depth_guided_selfcal" / "run.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "m7_depth_guided_selfcal" / "core.py",
            PROJECT_ROOT / "mm5_calib_benchmark" / "methods" / "alignment.py",
        ],
        "family": "深度引导自监督外参微调 + 深度投影 + boundary snap",
        "objective": "在不复用 MAR 历史 mask 的前提下，利用 RGB + raw depth + thermal 构建一条更强的 depth-guided thermal calibration 流程。",
        "implementation_note": "这是当前工作区的新算法，不是外部论文的一比一复现，而是融合了 Su 2025、depth bridge 和 T-ICP 思路的内部 composite 方法。",
        "workflow": [
            "从 train split 按类别抽样 48 个 scene，尝试做全局 pose 自监督搜索。",
            "对每个候选外参综合 edge score、MI、projected ratio 和 reprojection score 打分。",
            "在推理阶段使用 raw depth 做 3D 投影、z-buffer 和 support fill。",
            "做 scene-level similarity tune，再在 band 内执行 boundary snap refine。",
            "若 projected ratio 或 boundary guard 不满足条件，则按阶段回退。",
        ],
        "strengths": [
            "是当前 thermal benchmark 中表现最强的新方法。",
            "debug 信息完整，便于追踪 pose、自监督分数、投影率和边界改变量。",
        ],
        "limitations": [
            "当前缓存结果显示全局 pose refine 并没有通过质量门，因此最终保留的是 base pose + 局部 refinement。",
            "依赖 depth 图质量；若深度稀疏，收益会受限。",
        ],
        "references": [
            ("Su et al. 2025: cross-modal calibration", "https://doi.org/10.5194/isprs-annals-X-1-W2-2025-139-2025"),
            ("Muhovič & Perš 2023: parallax-aware multimodal mapping", "https://doi.org/10.3390/s23125676"),
            ("Cao et al. 2018: thermal-guided ICP for depth + thermal fusion", "https://doi.org/10.1364/OE.26.008179"),
            ("MM5 数据集论文", "https://doi.org/10.1016/j.inffus.2025.103516"),
        ],
    },
]


def _common_usage(method_key: str, track: str) -> str:
    return "\n".join(
        [
            "& .\\.venv\\Scripts\\python.exe -c \"from mm5_calib_benchmark.config import load_config; from mm5_calib_benchmark.pipeline import run_method; run_method(load_config(), "
            f"'{method_key}', '{track}')\"",
        ]
    )


def _full_suite_usage() -> str:
    return "& .\\.venv\\Scripts\\python.exe -m mm5_calib_benchmark.scripts.run_all_methods"


def _legacy_usage() -> str:
    return "\n".join(
        [
            "& .\\.venv\\Scripts\\python.exe -m mm5_calib_benchmark.scripts.run_legacy_mar_scene282 --mode engineered_best",
            "& .\\.venv\\Scripts\\python.exe -m mm5_calib_benchmark.scripts.run_legacy_mar_scene282 --mode paper_final",
        ]
    )


def _build_method_rows() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for meta in METHODS:
        summary = _summary_row(meta["output_name"])
        row = dict(meta)
        row["summary"] = summary
        rows.append(row)
    return rows


def _top_summary_rows(method_rows: list[dict[str, Any]]) -> list[list[str]]:
    sortable = [
        (
            float(method["summary"].get("mean_iou", 0.0)),
            method,
        )
        for method in method_rows
    ]
    sortable.sort(key=lambda item: item[0], reverse=True)
    return [
        [
            method["display"],
            method["track"],
            ", ".join(method["available_tracks"]),
            _fmt_float(method["summary"].get("mean_iou")),
            _fmt_float(method["summary"].get("pixel_accuracy")),
            _fmt_float(method["summary"].get("boundary_f1")),
            _fmt_float(method["summary"].get("valid_warp_coverage")),
            str(method["summary"].get("num_test_scenes", "N/A")),
        ]
        for _, method in sortable
    ]


def _split_counts() -> dict[str, int]:
    rows = _read_csv(OUTPUT_ROOT / "splits" / "index_with_splits.csv")
    counts: dict[str, int] = {}
    for row in rows:
        split = row.get("split", "").strip()
        counts[split] = counts.get(split, 0) + 1
    return counts


def _latest_scene_metrics_rows() -> list[list[str]]:
    rows = _read_csv(OUTPUT_ROOT / "scene_282_3_comparison" / "scene_282_3_metrics.csv")
    formatted: list[list[str]] = []
    for row in rows:
        formatted.append(
            [
                row.get("track", ""),
                row.get("method", ""),
                _fmt_float(row.get("mean_iou")),
                _fmt_float(row.get("pixel_accuracy")),
                _fmt_float(row.get("boundary_f1")),
                _fmt_float(row.get("valid_warp_coverage")),
            ]
        )
    return formatted


def _legacy_mar_metrics() -> dict[str, Any]:
    return _read_json(PROJECT_ROOT / "runs" / "manual_paper_engineered_metrics.json")


def _m7_cache() -> dict[str, Any]:
    return _read_json(
        OUTPUT_ROOT / "method_M7_depth_guided_selfcal_thermal" / "calib" / "pose_refine_cache.json"
    )


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def build_document() -> Path:
    RUNS_ROOT.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document)

    method_rows = _build_method_rows()
    split_counts = _split_counts()
    legacy_metrics = _legacy_mar_metrics()
    m7_cache = _m7_cache()

    title = document.add_heading("MM5 校准算法总结", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | 工作区：{_safe_rel(PROJECT_ROOT)}"
    )

    document.add_heading("1. 报告范围与当前评测设置", level=1)
    _add_bullets(
        document,
        [
            "本文总结的是当前工作区 benchmark 里已经落地并可运行的算法实现，而不只是对应论文名称。",
            "当前主对比图已经是 thermal-only 版本，主图中不再保留 UV 方向对比。",
            f"当前 split 文件包含 train={split_counts.get('train', 0)}、val={split_counts.get('val', 0)}、test={split_counts.get('test', 0)} 个 scene。",
            "当前 benchmark 运行时最多评估 30 个 test scenes（`runtime.max_test_scenes = 30`）。",
            "M3 目前在当前工作区里只有 UV benchmark 输出，因此单独说明，不放入 thermal 主图比较结论中。",
            "Legacy full MAR 结果作为历史参考保留，但它们来自同级 `MAR_test/backup_2600.py` 的历史 acceptance 输出，不是由新 benchmark 主路径重新计算得到。",
        ],
    )

    document.add_heading("2. 最新主对比图", level=1)
    _add_image(
        document,
        OUTPUT_ROOT / "scene_282_3_comparison" / "scene_282_3_thermal_only_comparison.png",
        "Scene 282_3 最新 thermal-only 对比图。",
        width_in=6.6,
    )

    document.add_heading("3. 当前 benchmark 总体排序", level=1)
    _add_table(
        document,
        ["方法", "主评测方向", "已支持方向", "mean_iou", "pixel_accuracy", "boundary_f1", "coverage", "test scenes"],
        _top_summary_rows(method_rows),
    )

    document.add_paragraph(
        "说明：上表把 thermal 与 UV 的主输出放在一起，只是为了完整汇总当前工作区的方法集合。"
        "如果只看最新主图和当前主线任务，应以 thermal 方法为主要比较对象。"
    )

    document.add_heading("4. 最新 Scene 282_3 面板指标", level=1)
    _add_table(
        document,
        ["track", "method", "mean_iou", "pixel_accuracy", "boundary_f1", "coverage"],
        _latest_scene_metrics_rows(),
    )

    document.add_heading("5. 各算法逐项总结", level=1)
    for index, method in enumerate(method_rows, start=1):
        document.add_heading(f"5.{index} {method['display']}", level=2)
        document.add_paragraph(f"主评测方向：{method['track']}")
        document.add_paragraph(f"当前工作区已支持方向：{', '.join(method['available_tracks'])}")
        document.add_paragraph(f"算法类型：{method['family']}")
        document.add_paragraph(f"设计目标：{method['objective']}")
        document.add_paragraph(f"当前实现说明：{method['implementation_note']}")

        document.add_paragraph("最新 benchmark 汇总指标：")
        _add_bullets(
            document,
            [
                f"mean_iou = {_fmt_float(method['summary'].get('mean_iou'))}",
                f"pixel_accuracy = {_fmt_float(method['summary'].get('pixel_accuracy'))}",
                f"boundary_f1 = {_fmt_float(method['summary'].get('boundary_f1'))}",
                f"valid_warp_coverage = {_fmt_float(method['summary'].get('valid_warp_coverage'))}",
                f"num_test_scenes = {method['summary'].get('num_test_scenes', 'N/A')}",
            ],
        )

        document.add_paragraph("当前流程逻辑：")
        _add_bullets(document, method["workflow"])

        document.add_paragraph("当前观察到的优势：")
        _add_bullets(document, method["strengths"])

        document.add_paragraph("当前观察到的局限：")
        _add_bullets(document, method["limitations"])

        document.add_paragraph("当前代码位置：")
        _add_bullets(document, [_safe_rel(path) for path in method["code_paths"]])

        document.add_paragraph("代表性输出图：")
        _add_image(
            document,
            method["image_path"],
            f"{method['display']} 代表性 qualitative grid（worst / median / best）。",
        )

        document.add_paragraph("单独运行该方法：")
        _add_code_block(document, _common_usage(method["key"], method["track"]))

        document.add_paragraph("运行完整 benchmark 套件：")
        _add_code_block(document, _full_suite_usage())

        document.add_paragraph("参考文献与网址：")
        for label, url in method["references"]:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{label}: ")
            _add_hyperlink(paragraph, url, url)

        if method["key"] == "m7":
            document.add_paragraph("M7 额外缓存观察：")
            _add_bullets(
                document,
                [
                    f"train_refine_scene_count = {m7_cache.get('train_scene_count', 'N/A')}",
                    f"score_before = {_fmt_float(m7_cache.get('score_before'))}",
                    f"score_after = {_fmt_float(m7_cache.get('score_after'))}",
                    f"score_gain = {_fmt_float(m7_cache.get('score_gain'))}",
                    f"used_refined_pose = {bool(m7_cache.get('used_refined_pose', False))}",
                    "当前缓存显示，全局 pose refine 没有通过质量门，因此最终结果保留的是 base pose，而不是更新后的全局外参。",
                ],
            )

        if index != len(method_rows):
            document.add_paragraph()

    document.add_page_break()
    document.add_heading("6. Legacy MAR 完整流程参考", level=1)
    document.add_paragraph(
        "这一节刻意与 M6 分开。M6 是 benchmark 内部重建的一条 MAR-style approximation，"
        "而下面这部分内容对应的是从同级 `MAR_test/backup_2600.py` 导入的历史完整 MAR full pipeline 结果。"
    )
    _add_image(
        document,
        OUTPUT_ROOT / "scene_282_3_comparison" / "scene_282_3_mar_history_panel.png",
        "Scene 282_3 导入的历史 MAR 对比面板。",
        width_in=6.2,
    )
    _add_bullets(
        document,
        [
            "MAR-engineered-full 在 raw thermal GT 上："
            f"pixel_accuracy={_fmt_float(0.9980712890625, 6)}, mean_iou={_fmt_float(0.991409943936322, 6)}",
            "MAR-paper-full 在 aligned GT 上："
            f"pixel_accuracy={_fmt_float(0.9935807291666666, 6)}, mean_iou={_fmt_float(0.9962712991363076, 6)}",
            "手工二值参考值与当前多类 benchmark 表格不是同一套指标定义。",
            "Legacy MAR 仍然是当前工作区里 Scene 282 最强的历史参考线，但如果不统一评测口径，不能把这些数值直接和新的 30-scene test 汇总结果等同起来。",
        ],
    )
    document.add_paragraph("来自 `runs/manual_paper_engineered_metrics.json` 的手工二值历史值：")
    _add_bullets(
        document,
        [
            "engineered_best 对 manual："
            f"pixel_accuracy_total={_fmt_float(legacy_metrics['engineered_best']['vs_manual']['pixel_accuracy_total'], 6)}, "
            f"foreground_iou={_fmt_float(legacy_metrics['engineered_best']['vs_manual']['foreground_iou'], 6)}",
            "paper_final 对 manual："
            f"pixel_accuracy_total={_fmt_float(legacy_metrics['paper_final']['vs_manual']['pixel_accuracy_total'], 6)}, "
            f"foreground_iou={_fmt_float(legacy_metrics['paper_final']['vs_manual']['foreground_iou'], 6)}",
        ],
    )
    document.add_paragraph("复现 Legacy Scene 282 结果的运行方式：")
    _add_code_block(document, _legacy_usage())
    document.add_paragraph("公开参考网址：")
    ref_par = document.add_paragraph(style="List Bullet")
    ref_par.add_run("包含 MM5 与公开 MAR 背景的论文：")
    _add_hyperlink(ref_par, "https://doi.org/10.1016/j.inffus.2025.103516", "https://doi.org/10.1016/j.inffus.2025.103516")

    document.add_page_break()
    document.add_heading("7. 关键输出路径", level=1)
    _add_bullets(
        document,
        [
            f"最新主对比图：{_safe_rel(OUTPUT_ROOT / 'scene_282_3_comparison' / 'scene_282_3_thermal_only_comparison.png')}",
            f"最新 scene 指标表：{_safe_rel(OUTPUT_ROOT / 'scene_282_3_comparison' / 'scene_282_3_metrics.csv')}",
            f"M7 汇总 JSON：{_safe_rel(OUTPUT_ROOT / 'method_M7_depth_guided_selfcal_thermal' / 'metrics' / 'summary.json')}",
            f"M7 pose 缓存：{_safe_rel(OUTPUT_ROOT / 'method_M7_depth_guided_selfcal_thermal' / 'calib' / 'pose_refine_cache.json')}",
            f"Legacy MAR 备注：{_safe_rel(OUTPUT_ROOT / 'scene_282_3_comparison' / 'scene_282_3_mar_history_note.md')}",
        ],
    )

    document.add_heading("8. 简短结论", level=1)
    _add_bullets(
        document,
        [
            "在当前启用的 thermal 方法中，M7 是 30-scene test 汇总结果里最强的新 benchmark 方法。",
            "在当前 MM5 thermal 设定下，纯 homography 家族（M0-M2）明显弱于有 depth 支撑的方法（M4-M7）。",
            "M6 应被理解为 benchmark 侧的 MAR-style refine 实验，而不能等同于旧版完整 MAR full pipeline。",
            "历史 full MAR 仍然是 Scene 282 的高参考线，但由于评测协议不同，导入的 acceptance 数值需要谨慎解释。",
        ],
    )

    document.save(OUTPUT_DOC)
    return OUTPUT_DOC


def main() -> None:
    output_path = build_document()
    print(output_path)


if __name__ == "__main__":
    main()
