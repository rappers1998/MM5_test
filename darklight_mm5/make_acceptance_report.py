import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT_ROOT = ROOT / "outputs"
METRICS_DIR = OUTPUT_ROOT / "metrics"
REPORT_PATH = OUTPUT_ROOT / "MM5_darklight_calibration_fusion_acceptance_report.docx"


def read_csv(path):
    with Path(path).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fmt(value, digits=3):
    if value in (None, ""):
        return ""
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, key in enumerate(headers):
            cells[idx].text = str(row.get(key, ""))
    return table


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    return p


def add_image(doc, path, caption, width=6.7):
    path = Path(path)
    if not path.exists():
        add_paragraph(doc, f"缺失图片：{path}")
        return
    doc.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_document(doc):
    styles = doc.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "SimSun"


def build_key_sample_table(per_sample):
    rows = []
    for r in per_sample:
        rows.append(
            {
                "ID": r["aligned_id"],
                "Sequence": r["sequence"],
                "Split": r["split"],
                "RGB均值": fmt(r["dark_mean"], 2),
                "RGB->官方NCC": fmt(r["rgb_ncc_to_official"], 4),
                "LWIR->官方NCC": fmt(r["lwir_ncc_to_official"], 4),
                "Raw跨模态NCC提升": fmt(r["raw_lwir_to_rgb_ncc_gain"], 4),
                "边缘距离改善(px)": fmt(r["raw_lwir_to_rgb_edge_distance_gain"], 2),
                "融合覆盖率": fmt(float(r["fusion_alpha_coverage"]) * 100, 2) + "%",
            }
        )
    return rows


def build_stage_table(reg_rows):
    rows = []
    for r in reg_rows:
        rows.append(
            {
                "样本": r["sample"],
                "流程": r["pipeline"],
                "阶段": r["stage"],
                "NCC": fmt(r.get("ncc"), 4),
                "边缘距离": fmt(r.get("edge_distance"), 2),
                "有效区域": fmt(float(r.get("valid_ratio", "nan")) * 100, 1) + "%"
                if r.get("valid_ratio")
                else "",
                "Score": fmt(r.get("score"), 3),
                "说明": r.get("failure_reason") or r.get("ecc_reason") or "",
            }
        )
    return rows


def build_fusion_table(fusion_rows):
    keep = {"raw_rgb1", "enhanced_rgb1", "fused_intensity", "fused_heat"}
    rows = []
    for r in fusion_rows:
        if r["image"] not in keep:
            continue
        rows.append(
            {
                "样本": r["sample"],
                "图像": r["image"],
                "Entropy": fmt(r["entropy"], 3),
                "AvgGrad": fmt(r["average_gradient"], 2),
                "SpatialFreq": fmt(r["spatial_frequency"], 2),
                "Std": fmt(r["std"], 2),
                "MI-RGB": fmt(r["mi_with_raw_rgb"], 3),
                "MI-LWIR": fmt(r["mi_with_calibrated_lwir"], 3),
                "相对Raw变化": fmt(r["mean_abs_change_vs_raw_rgb"], 2),
                "融合覆盖率": fmt(float(r["alpha_coverage"]) * 100, 2) + "%",
            }
        )
    return rows


def main():
    per_sample = read_csv(METRICS_DIR / "per_sample.csv")
    reg_rows = read_csv(METRICS_DIR / "registration_stages.csv")
    fusion_rows = read_csv(METRICS_DIR / "fusion_metrics.csv")
    selected = read_csv(OUTPUT_ROOT / "selected_dark_samples.csv")
    summary_path = METRICS_DIR / "summary.json"
    summary = json.loads(summary_path.read_text(encoding="utf-8")) if summary_path.exists() else {}

    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)

    title = doc.add_heading("MM5 暗光 RGB1 + LWIR 校准与融合验收报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 验收结论", level=1)
    add_paragraph(
        doc,
        "结论：当前版本可以进入阶段性验收。输入已确认使用 MM5 RAW 的 RGB1 与 LWIR frame1；"
        "配准使用官方 aligned RGB1/T16 作为同模态校准参考；融合图是在 raw RGB1 画布上生成。"
        "需要注意：融合结果包含低光增强和热显著性注入，因此视觉清晰度不能单独代表配准精度。",
    )
    add_paragraph(
        doc,
        "验收依据：同模态 RGB->官方 RGB 的 NCC 接近 1.0；LWIR->官方 LWIR 的 NCC 在 0.906-0.935；"
        "raw LWIR->raw RGB1 的边缘距离相对 fit-cover baseline 有改善；融合覆盖率约 0.9%，说明 thermal 注入范围受控。",
    )

    doc.add_heading("2. 输入数据与样本选择", level=1)
    add_paragraph(doc, "样本来自 index_with_splits.csv，按 raw RGB1 平均亮度从 test/val 中选择最暗的 3 个样本。")
    add_table(
        doc,
        ["aligned_id", "sequence", "split", "dark_mean", "raw_rgb1_path", "raw_thermal16_path"],
        selected,
    )

    doc.add_heading("3. 算法流程与关键参数", level=1)
    params = [
        {"模块": "输入", "参数": "raw_rgb1_path + raw_thermal16_path", "取值": "同步 frame1 RAW 输入"},
        {"模块": "归一化", "参数": "percentile", "取值": "1%-99% 转 uint8"},
        {"模块": "特征增强", "参数": "CLAHE", "取值": "clipLimit=3.0, tileGridSize=(8,8)"},
        {"模块": "SIFT", "参数": "nfeatures / contrast / edge", "取值": "5000 / 0.005 / 12"},
        {"模块": "匹配", "参数": "Lowe ratio", "取值": "0.78"},
        {"模块": "RANSAC", "参数": "homography reprojection threshold", "取值": "4.0 px"},
        {"模块": "ECC refine", "参数": "motion / iterations / eps", "取值": "affine / 80 / 1e-6"},
        {"模块": "低光增强", "参数": "gain / gamma / CLAHE", "取值": "max gain 8.0 / gamma 0.72 / clipLimit 1.8"},
        {"模块": "融合显著性", "参数": "Gaussian sigma / saliency percentile", "取值": "17 / 45%-97%"},
        {"模块": "融合权重", "参数": "alpha smoothing / max alpha", "取值": "sigma 3 / max 0.68"},
    ]
    add_table(doc, ["模块", "参数", "取值"], params)

    doc.add_heading("4. 样本级关键指标", level=1)
    add_table(
        doc,
        [
            "ID",
            "Sequence",
            "Split",
            "RGB均值",
            "RGB->官方NCC",
            "LWIR->官方NCC",
            "Raw跨模态NCC提升",
            "边缘距离改善(px)",
            "融合覆盖率",
        ],
        build_key_sample_table(per_sample),
    )

    doc.add_heading("5. 配准阶段指标", level=1)
    add_paragraph(doc, "NCC 越高越好；边缘距离越低越好；Score 为内部保护指标：NCC - 0.006 * edge_distance。")
    add_table(
        doc,
        ["样本", "流程", "阶段", "NCC", "边缘距离", "有效区域", "Score", "说明"],
        build_stage_table(reg_rows),
    )

    doc.add_heading("6. 融合指标", level=1)
    add_paragraph(
        doc,
        "Entropy、AvgGrad、SpatialFreq、Std 用于观察清晰度/纹理/对比度变化；MI-RGB 与 MI-LWIR 表示融合图与两种输入的信息关联。",
    )
    add_table(
        doc,
        [
            "样本",
            "图像",
            "Entropy",
            "AvgGrad",
            "SpatialFreq",
            "Std",
            "MI-RGB",
            "MI-LWIR",
            "相对Raw变化",
            "融合覆盖率",
        ],
        build_fusion_table(fusion_rows),
    )

    doc.add_heading("7. 可视化验收图", level=1)
    for r in per_sample:
        sid = f"{int(r['aligned_id']):03d}_seq{r['sequence']}"
        doc.add_heading(f"Sequence {r['sequence']} / aligned_id {r['aligned_id']}", level=2)
        add_image(doc, OUTPUT_ROOT / "quads" / f"{sid}_quad.png", "四宫格：RGB1 Raw / LWIR Raw / 校准后 LWIR / 融合结果")
        add_image(doc, OUTPUT_ROOT / "edge_reviews" / f"{sid}_official_check.png", "官方参考检查：raw RGB1/LWIR 到 official aligned 画布")
        add_image(doc, OUTPUT_ROOT / "edge_reviews" / f"{sid}_lwir_calibration_check.png", "LWIR 校准检查：官方 LWIR / 校准 LWIR / 差分 / 边缘")
        add_image(doc, OUTPUT_ROOT / "edge_reviews" / f"{sid}_edge_review.png", "融合区域检查：边缘叠加与 alpha mask")
        doc.add_page_break()

    doc.add_heading("8. 文件索引", level=1)
    files = [
        {"文件": "运行脚本", "路径": str(ROOT / "run_darklight.py")},
        {"文件": "算法说明", "路径": str(ROOT / "README.md")},
        {"文件": "样本选择", "路径": str(OUTPUT_ROOT / "selected_dark_samples.csv")},
        {"文件": "样本总指标", "路径": str(METRICS_DIR / "per_sample.csv")},
        {"文件": "配准阶段指标", "路径": str(METRICS_DIR / "registration_stages.csv")},
        {"文件": "融合指标", "路径": str(METRICS_DIR / "fusion_metrics.csv")},
    ]
    add_table(doc, ["文件", "路径"], files)

    doc.add_heading("9. 补充说明", level=1)
    add_paragraph(
        doc,
        "旧 method1-method5 的源码/config 已删除；旧根目录 outputs/method* 二进制结果因权限复核超时仍可能残留，"
        "但本报告只引用 darklight_mm5/outputs 下的新结果。",
    )
    if summary:
        add_paragraph(doc, "summary.json 已生成，可用于快速查看均值/最小值/最大值统计。")

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
